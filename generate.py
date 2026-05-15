#!/usr/bin/env python3
import os, sys, json, datetime, urllib.request, urllib.error, logging, hashlib
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    PDF_OK = True
except ImportError:
    PDF_OK = False

BASE_DIR        = Path(__file__).parent
DATA_MD         = BASE_DIR / "data.md"
OLD_RESUMES_DIR = BASE_DIR / "old_resumes"
OUTPUT_DIR      = BASE_DIR / "output"
CACHE_DIR       = BASE_DIR / "cache"
MODEL           = "inclusionai/ling-2.6-1t:free"
OPENROUTER_URL  = "https://openrouter.ai/api/v1/chat/completions"
OUTPUT_DIR.mkdir(exist_ok=True)
CACHE_DIR.mkdir(exist_ok=True)

# ── language config ────────────────────────────────────────────────────────────

LANGUAGES = {
    "1": {
        "name": "Portuguese (Brazil)",
        "resume_headers": {
            "name":       "NOME",
            "contact":    "CONTATO",
            "summary":    "RESUMO PROFISSIONAL",
            "experience": "EXPERIÊNCIA PROFISSIONAL",
            "education":  "FORMAÇÃO ACADÊMICA",
            "skills":     "HABILIDADES",
        },
    },
    "2": {
        "name": "English",
        "resume_headers": {
            "name":       "NAME",
            "contact":    "CONTACT",
            "summary":    "PROFESSIONAL SUMMARY",
            "experience": "WORK EXPERIENCE",
            "education":  "EDUCATION",
            "skills":     "SKILLS",
        },
    },
    "3": {
        "name": "Spanish",
        "resume_headers": {
            "name":       "NOMBRE",
            "contact":    "CONTACTO",
            "summary":    "RESUMEN PROFESIONAL",
            "experience": "EXPERIENCIA PROFESIONAL",
            "education":  "FORMACIÓN ACADÉMICA",
            "skills":     "HABILIDADES",
        },
    },
}

# ── readers ────────────────────────────────────────────────────────────────────

def read_data_md():
    if not DATA_MD.exists():
        return ""
    return DATA_MD.read_text(encoding="utf-8")

def read_old_resumes():
    if not OLD_RESUMES_DIR.exists():
        return ""
    parts = []
    for file in sorted(OLD_RESUMES_DIR.iterdir()):
        if file.suffix.lower() not in {".md", ".txt", ".pdf", ".docx"}:
            continue
        content = ""
        if file.suffix.lower() in {".md", ".txt"}:
            content = file.read_text(encoding="utf-8", errors="ignore")
        elif file.suffix.lower() == ".pdf":
            content = _read_pdf(file)
        elif file.suffix.lower() == ".docx":
            content = _read_docx(file)
        if content.strip():
            parts.append(f"[Resume: {file.name}]\n{content}")
    return "\n\n---\n\n".join(parts)

def _read_pdf(path):
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    except ImportError:
        pass
    return ""

def _read_docx(path):
    if not DOCX_OK:
        return ""
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)

# ── prompt ─────────────────────────────────────────────────────────────────────

def build_prompt(raw_data, resumes, job_description, lang):
    h = lang["resume_headers"]
    lang_name = lang["name"]
    section_resumes = f"\n=== OLD RESUMES ===\n{resumes}\n" if resumes.strip() else ""

    return f"""You are an expert in recruitment and ATS (Applicant Tracking System) resume optimization.

Candidate raw data (may be disorganized, from multiple sources):

=== RAW DATA ===
{raw_data}
{section_resumes}
=== JOB DESCRIPTION ===
{job_description}

=== CRITICAL LANGUAGE RULE ===
Output language: {lang_name}
Write EVERY word in {lang_name} — section headers, job titles, bullet points, summary, skills, everything.
If raw data is in another language, translate ALL content to {lang_name}.
Never mix languages.

=== OUTPUT FORMAT — FOLLOW EXACTLY ===
{h["name"]}: [full name]
{h["contact"]}: [email] | [phone] | [linkedin] | [city/state]

## {h["summary"]}
[3-4 lines in {lang_name}, tailored to the job]

## {h["experience"]}

### [Job title in {lang_name}] - [Company] | [Period]
- [achievement in {lang_name}]
- [achievement in {lang_name}]

## {h["education"]}

### [Degree in {lang_name}] - [Institution] | [Year]

## {h["skills"]}
[comma-separated list in {lang_name}]

=== RULES ===
1. Extract name, contact, experiences, education, skills, achievements
2. Ignore duplicates and irrelevant items for this job
3. Use the SAME keywords from the job description (crucial for ATS)
4. Quantify achievements with numbers whenever possible
5. Max 1-2 pages
6. Reply ONLY with the resume — no explanations, no code fences, no commentary

=== SKILLS SECTION — STRICT RULES ===
List ONLY concrete, specific, verifiable technical skills:
  ✅ INCLUDE: programming languages, frameworks, databases, tools, platforms, certifications
     Examples: Python, Django, PostgreSQL, Docker, AWS, Git, React, Kubernetes

  ❌ NEVER include any of the following — they are obvious, generic or not real skills:
     - Soft skills: "teamwork", "communication", "leadership", "proactivity", "organization"
     - Generic responsibilities: "code review", "software architecture", "third-party integrations",
       "data security", "compliance", "agile methodology", "scrum"
     - Languages spoken: "English", "Portuguese", "Spanish", "native", "intermediate", "fluent"
       (spoken languages do NOT belong in the skills section)
     - Vague tools: "Microsoft Office", "Google Docs", "email", "Slack"

If a skill could appear on ANY professional's resume regardless of their technical area, do NOT include it.
"""

# ── validation ─────────────────────────────────────────────────────────────────

def validate_job_description(text: str) -> tuple:
    """Returns (is_valid, error_key) where error_key maps to i18n T dict."""
    stripped = (text or "").strip()
    if not stripped:
        return False, "no_job_error"
    if len(stripped) < 30:
        return False, "job_too_short"
    return True, ""

# ── cache ───────────────────────────────────────────────────────────────────────

def _cache_key(prompt: str, model: str) -> str:
    return hashlib.sha256(f"{model}:{prompt}".encode()).hexdigest()

def _cache_get(key: str):
    path = CACHE_DIR / f"{key}.json"
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))["response"]
    except Exception:
        return None

def _cache_set(key: str, response: str) -> None:
    path = CACHE_DIR / f"{key}.json"
    path.write_text(
        json.dumps({"response": response, "created_at": datetime.datetime.now().isoformat()}),
        encoding="utf-8",
    )

# ── API ────────────────────────────────────────────────────────────────────────

def call_model(prompt, model):
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        logger.error("OPENROUTER_API_KEY is not set")
        sys.exit(1)

    key = _cache_key(prompt, model)
    cached = _cache_get(key)
    if cached:
        logger.info("Cache hit — skipping API call (model=%s)", model)
        return cached

    logger.info("Calling OpenRouter (model=%s)", model)
    payload = json.dumps({
        "model": model,
        "max_tokens": 4096,
        "messages": [{"role": "user", "content": prompt}]
    }).encode("utf-8")

    req = urllib.request.Request(
        OPENROUTER_URL, data=payload,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://resume-ats-local",
            "X-Title": "ATS Resume Generator",
        }, method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8")
        logger.error("OpenRouter HTTP %s: %s", e.code, body)
        sys.exit(1)

    try:
        content = data["choices"][0]["message"]["content"]
        if not content:
            logger.error("Model returned empty response (model=%s)", model)
            sys.exit(1)
        content = content.strip()
        if content.startswith("```"):
            lines = content.splitlines()
            lines = [l for l in lines if not l.strip().startswith("```")]
            content = "\n".join(lines).strip()
        _cache_set(key, content)
        logger.info("Response cached (model=%s, chars=%d)", model, len(content))
        return content
    except (KeyError, IndexError):
        logger.error("Unexpected API response structure: %s", json.dumps(data))
        sys.exit(1)

# ── savers ─────────────────────────────────────────────────────────────────────

def save_docx(text, base_name, lang):
    if not DOCX_OK:
        logger.warning("python-docx not installed; DOCX output skipped")
        return None
    h = lang["resume_headers"]
    name_lbl    = h["name"]
    contact_lbl = h["contact"]

    path = OUTPUT_DIR / f"{base_name}.docx"
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    for line in text.strip().splitlines():
        line = line.rstrip()
        if line.startswith(f"{name_lbl}:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[len(name_lbl)+1:].strip())
            run.bold = True; run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        elif line.startswith(f"{contact_lbl}:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[len(contact_lbl)+1:].strip())
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            doc.add_paragraph()
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip().upper())
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            p2 = doc.add_paragraph()
            run2 = p2.add_run("─" * 60)
            run2.font.size = Pt(7)
            run2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            p2.paragraph_format.space_after = Pt(4)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(1)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(9.5)
        elif not line:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(2)

    doc.save(str(path))
    return path

def save_pdf(text, base_name, lang):
    if not PDF_OK:
        logger.warning("reportlab not installed; PDF output skipped")
        return None
    h = lang["resume_headers"]
    name_lbl    = h["name"]
    contact_lbl = h["contact"]

    path = OUTPUT_DIR / f"{base_name}.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    color_p = colors.HexColor("#000000")
    color_t = colors.HexColor("#000000")
    color_s = colors.HexColor("#444444")
    s_name    = ParagraphStyle("Name",    parent=styles["Title"],    fontSize=20, textColor=color_t, spaceAfter=2,   alignment=1, fontName="Helvetica-Bold")
    s_contact = ParagraphStyle("Contact", parent=styles["Normal"],   fontSize=9,  textColor=color_s, spaceAfter=12,  alignment=1, fontName="Helvetica")
    s_section = ParagraphStyle("Section", parent=styles["Heading2"], fontSize=10, textColor=color_p, spaceBefore=12, spaceAfter=2, fontName="Helvetica-Bold")
    s_sub     = ParagraphStyle("Sub",     parent=styles["Normal"],   fontSize=10, textColor=color_t, spaceBefore=6,  spaceAfter=1, fontName="Helvetica-Bold")
    s_bullet  = ParagraphStyle("Bullet",  parent=styles["Normal"],   fontSize=9.5,textColor=color_t, leftIndent=12,  spaceAfter=2, fontName="Helvetica")
    s_normal  = ParagraphStyle("Normal2", parent=styles["Normal"],   fontSize=9.5,textColor=color_t, spaceAfter=2,   fontName="Helvetica")

    story = []
    for line in text.strip().splitlines():
        line = line.rstrip()
        if line.startswith(f"{name_lbl}:"):
            story.append(Paragraph(line[len(name_lbl)+1:].strip(), s_name))
        elif line.startswith(f"{contact_lbl}:"):
            story.append(Paragraph(line[len(contact_lbl)+1:].strip(), s_contact))
        elif line.startswith("## "):
            story.append(HRFlowable(width="100%", thickness=1, color=color_p, spaceAfter=2))
            story.append(Paragraph(line[3:].strip(), s_section))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:].strip(), s_sub))
        elif line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph("- " + line[2:].strip(), s_bullet))
        elif not line:
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(line, s_normal))

    doc.build(story)
    return path

