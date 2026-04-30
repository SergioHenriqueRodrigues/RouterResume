#!/usr/bin/env python3
import os, sys, json, argparse, datetime, textwrap, urllib.request, urllib.error
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    PDF_OK = True
except ImportError:
    PDF_OK = False

BASE_DIR       = Path(__file__).parent
DADOS_MD       = BASE_DIR / "dados.md"
CURRICULOS_DIR = BASE_DIR / "curriculos_antigos"
OUTPUT_DIR     = BASE_DIR / "output"
MODEL          = "inclusionai/ling-2.6-1t:free"
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
OUTPUT_DIR.mkdir(exist_ok=True)

def ler_dados_md():
    if not DADOS_MD.exists():
        print("AVISO: dados.md nao encontrado.")
        return ""
    return DADOS_MD.read_text(encoding="utf-8")

def ler_curriculos_antigos():
    if not CURRICULOS_DIR.exists():
        return ""
    partes = []
    for arquivo in sorted(CURRICULOS_DIR.iterdir()):
        if arquivo.suffix.lower() not in {".md", ".txt", ".pdf", ".docx"}:
            continue
        conteudo = ""
        if arquivo.suffix.lower() in {".md", ".txt"}:
            conteudo = arquivo.read_text(encoding="utf-8", errors="ignore")
        elif arquivo.suffix.lower() == ".pdf":
            conteudo = _ler_pdf(arquivo)
        elif arquivo.suffix.lower() == ".docx":
            conteudo = _ler_docx(arquivo)
        if conteudo.strip():
            partes.append(f"[Curriculo: {arquivo.name}]\n{conteudo}")
    return "\n\n---\n\n".join(partes)

def _ler_pdf(path):
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    except ImportError:
        pass
    print(f"  AVISO: instale pdfplumber para ler {path.name}")
    return ""

def _ler_docx(path):
    if not DOCX_OK:
        return ""
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)

def montar_prompt(dados_brutos, curriculos, vaga, idioma):
    secao_curriculos = f"\n=== CURRICULOS ANTIGOS ===\n{curriculos}\n" if curriculos.strip() else ""
    return f"""Voce e um especialista em recrutamento e otimizacao de curriculos ATS.

Dados brutos do candidato (podem estar desordenados):

=== DADOS BRUTOS ===
{dados_brutos}
{secao_curriculos}
=== VAGA ===
{vaga}

Idioma: {idioma}

Instrucoes:
1. Extraia nome, contato, experiencias, formacao, habilidades, conquistas
2. Ignore duplicatas e irrelevantes para essa vaga
3. Resumo profissional focado no que a vaga pede (3-4 linhas)
4. Use EXATAMENTE as palavras-chave da vaga (crucial para ATS)
5. Quantifique conquistas sempre que possivel
6. Maximo 1-2 paginas

Responda SOMENTE com o curriculo formatado, sem explicacoes, sem blocos de codigo, sem markdown extra.
Use este formato exato:

NOME: [nome completo]
CONTATO: [email] | [telefone] | [linkedin] | [cidade]

## RESUMO PROFISSIONAL
[paragrafo]

## EXPERIENCIA PROFISSIONAL

### [Cargo] - [Empresa] | [Periodo]
- [conquista]
- [conquista]

## FORMACAO ACADEMICA

### [Curso] - [Instituicao] | [Ano]

## HABILIDADES
[lista]
"""

def chamar_modelo(prompt, modelo):
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        print("ERRO: OPENROUTER_API_KEY nao definida.")
        print('  setx OPENROUTER_API_KEY "sk-or-..."')
        sys.exit(1)
    print(f"Chamando OpenRouter ({modelo})...")
    payload = json.dumps({
        "model": modelo,
        "max_tokens": 4096,
        "messages": [{"role": "user", "content": prompt}]
    }).encode("utf-8")
    req = urllib.request.Request(
        OPENROUTER_URL, data=payload,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://curriculo-ats-local",
            "X-Title": "Gerador Curriculo ATS",
        }, method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        print(f"ERRO na API ({e.code}): {e.read().decode('utf-8')}")
        sys.exit(1)
    try:
        content = data["choices"][0]["message"]["content"]
        if not content:
            print("ERRO: modelo retornou vazio. Tente outro --modelo")
            sys.exit(1)
        # Remove blocos de codigo se o modelo colocar
        content = content.strip()
        if content.startswith("```"):
            lines = content.splitlines()
            lines = [l for l in lines if not l.strip().startswith("```")]
            content = "\n".join(lines).strip()
        return content
    except (KeyError, IndexError):
        print(f"ERRO resposta inesperada: {json.dumps(data, indent=2)}")
        sys.exit(1)

def salvar_md(texto, nome_base):
    path = OUTPUT_DIR / f"{nome_base}.md"
    path.write_text(texto, encoding="utf-8")
    return path

def salvar_docx(texto, nome_base):
    if not DOCX_OK:
        print("  AVISO: rode: pip install python-docx")
        return None
    path = OUTPUT_DIR / f"{nome_base}.docx"
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
    for linha in texto.strip().splitlines():
        linha = linha.rstrip()
        if linha.startswith("NOME:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(linha.replace("NOME:", "").strip())
            run.bold = True; run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        elif linha.startswith("CONTATO:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(linha.replace("CONTATO:", "").strip())
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            doc.add_paragraph()
        elif linha.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(linha[3:].strip().upper())
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            p2 = doc.add_paragraph()
            run2 = p2.add_run("-" * 60)
            run2.font.size = Pt(7)
            run2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
            p2.paragraph_format.space_after = Pt(4)
        elif linha.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(linha[4:].strip())
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(1)
        elif linha.startswith("- ") or linha.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(linha[2:].strip())
            run.font.size = Pt(9.5)
        elif not linha:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph()
            run = p.add_run(linha)
            run.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(2)
    doc.save(str(path))
    return path

def salvar_pdf(texto, nome_base):
    if not PDF_OK:
        print("  AVISO: rode: pip install reportlab")
        return None
    path = OUTPUT_DIR / f"{nome_base}.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    cor_p = colors.HexColor("#2E86AB")
    cor_t = colors.HexColor("#1A1A2E")
    cor_s = colors.HexColor("#555555")
    s_nome    = ParagraphStyle("Nome",    parent=styles["Title"],   fontSize=20, textColor=cor_t, spaceAfter=2,   alignment=1, fontName="Helvetica-Bold")
    s_contato = ParagraphStyle("Contato", parent=styles["Normal"],  fontSize=9,  textColor=cor_s, spaceAfter=12,  alignment=1, fontName="Helvetica")
    s_secao   = ParagraphStyle("Secao",   parent=styles["Heading2"],fontSize=10, textColor=cor_p, spaceBefore=12, spaceAfter=2, fontName="Helvetica-Bold")
    s_sub     = ParagraphStyle("Sub",     parent=styles["Normal"],  fontSize=10, textColor=cor_t, spaceBefore=6,  spaceAfter=1, fontName="Helvetica-Bold")
    s_bullet  = ParagraphStyle("Bullet",  parent=styles["Normal"],  fontSize=9.5,textColor=cor_t, leftIndent=12,  spaceAfter=2, fontName="Helvetica")
    s_normal  = ParagraphStyle("Normal2", parent=styles["Normal"],  fontSize=9.5,textColor=cor_t, spaceAfter=2,   fontName="Helvetica")
    story = []
    for linha in texto.strip().splitlines():
        linha = linha.rstrip()
        if linha.startswith("NOME:"):
            story.append(Paragraph(linha.replace("NOME:", "").strip(), s_nome))
        elif linha.startswith("CONTATO:"):
            story.append(Paragraph(linha.replace("CONTATO:", "").strip(), s_contato))
        elif linha.startswith("## "):
            story.append(HRFlowable(width="100%", thickness=1, color=cor_p, spaceAfter=2))
            story.append(Paragraph(linha[3:].strip(), s_secao))
        elif linha.startswith("### "):
            story.append(Paragraph(linha[4:].strip(), s_sub))
        elif linha.startswith("- ") or linha.startswith("* "):
            story.append(Paragraph("- " + linha[2:].strip(), s_bullet))
        elif not linha:
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(linha, s_normal))
    doc.build(story)
    return path

def main():
    parser = argparse.ArgumentParser(description="Gerador de Curriculo ATS")
    parser.add_argument("--vaga",    help="Arquivo .txt com o texto da vaga")
    parser.add_argument("--idioma",  default="portugues do Brasil")
    parser.add_argument("--formato", choices=["md", "docx", "pdf", "todos"], default="docx")
    parser.add_argument("--modelo",  default=MODEL)
    args = parser.parse_args()

    print("\nGerador de Curriculo ATS")
    print("-" * 40)

    dados_brutos = ler_dados_md()
    print(f"Lendo {DADOS_MD.name}... OK")

    curriculos = ler_curriculos_antigos()
    n = curriculos.count("[Curriculo:")
    print(f"Curriculos antigos: {n} arquivo(s)")

    if args.vaga:
        vaga_path = Path(args.vaga)
        if not vaga_path.exists():
            print(f"ERRO: {args.vaga} nao encontrado")
            sys.exit(1)
        vaga = vaga_path.read_text(encoding="utf-8")
        print(f"Vaga: {args.vaga}")
    else:
        print("\nCole a vaga. Termine com Enter + Ctrl+Z (Windows):\n")
        linhas = []
        try:
            while True:
                linhas.append(input())
        except EOFError:
            pass
        vaga = "\n".join(linhas).strip()
        if not vaga:
            print("ERRO: nenhuma vaga fornecida.")
            sys.exit(1)

    prompt    = montar_prompt(dados_brutos, curriculos, vaga, args.idioma)
    curriculo = chamar_modelo(prompt, args.modelo)

    ts        = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    nome_base = f"curriculo_ats_{ts}"

    print(f"\nSalvando em output/...")

    if args.formato in ("md", "todos"):
        p = salvar_md(curriculo, nome_base)
        print(f"  OK: {p.name}")

    if args.formato in ("docx", "todos"):
        p = salvar_docx(curriculo, nome_base)
        if p: print(f"  OK: {p.name}")

    if args.formato in ("pdf", "todos"):
        p = salvar_pdf(curriculo, nome_base)
        if p: print(f"  OK: {p.name}")

    print(f"\nArquivos em: {OUTPUT_DIR}")
    print("-" * 40)
    print("\nPrevia:\n")
    print(textwrap.indent(curriculo[:600] + ("..." if len(curriculo) > 600 else ""), "  "))
    print()

if __name__ == "__main__":
    main()
